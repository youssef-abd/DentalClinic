"""
Invoice service for handling invoice generation and Word template automation
"""

import os
import json
from datetime import datetime, date
from typing import List, Dict, Optional, Tuple
from decimal import Decimal, ROUND_HALF_UP
import tempfile
import shutil

try:
    from docx import Document
    from docx.shared import Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.shared import OxmlElement, qn
    from docx.oxml.ns import nsdecls
    from docx.oxml import parse_xml
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    # Create dummy classes for type hints when docx is not available
    class Document:
        pass
    class Inches:
        pass
    class WD_ALIGN_PARAGRAPH:
        CENTER = None
        RIGHT = None
    print("Warning: python-docx not available. Install with: pip install python-docx")

from ..models.database import Visit, Patient, DatabaseManager

class InvoiceService:
    """Service for handling invoice operations and Word template automation"""
    
    def __init__(self, db_manager: DatabaseManager):
        self.db_manager = db_manager
        self.template_folder = os.path.expanduser("~/.dentistedb/templates")
        os.makedirs(self.template_folder, exist_ok=True)
        
        # Doctor information from the template
        self.doctor_info = {
            "name": "Dr. Mouna Afquir",
            "title": "Médecin Dentiste",
            "address": "68, rue: Patrice Lumumba, appt. 4 Rabat-10010",
            "phone": "0537 76 56 55",
            "inpe": "104003025",
            "if": "34204085",
            "taxe_professionnelle": "25121662",
            "cnss": "6912811",
            "ice": "001890220000028"
        }
    
    def get_visits_for_invoice(self, visit_ids: List[int]) -> List[Visit]:
        """Get visits by IDs for invoice generation"""
        session = self.db_manager.get_session()
        try:
            visits = session.query(Visit).filter(Visit.id.in_(visit_ids)).all()
            return visits
        finally:
            session.close()
    
    def get_patient_by_visit_ids(self, visit_ids: List[int]) -> Optional[Patient]:
        """Get patient information from visit IDs"""
        session = self.db_manager.get_session()
        try:
            visit = session.query(Visit).filter(Visit.id.in_(visit_ids)).first()
            if visit:
                return session.query(Patient).filter_by(id=visit.patient_id).first()
            return None
        finally:
            session.close()
    
    def calculate_invoice_totals(self, visits: List[Visit]) -> Dict[str, float]:
        """Calculate invoice totals from visits"""
        subtotal = sum(visit.prix for visit in visits if visit.prix)
        tax_rate = 0.20  # 20% TVA
        tax_amount = subtotal * tax_rate
        total = subtotal + tax_amount
        
        return {
            "subtotal": round(subtotal, 2),
            "tax_rate": tax_rate,
            "tax_amount": round(tax_amount, 2),
            "total": round(total, 2)
        }
    
    def aggregate_treatments(self, visits: List[Visit]) -> List[Dict]:
        """Aggregate treatments from visits into invoice line items"""
        treatment_groups = {}
        
        for visit in visits:
            if not visit.acte or not visit.prix:
                continue
                
            # Create a key for grouping similar treatments
            treatment_key = f"{visit.acte}_{visit.dent or 'N/A'}"
            
            if treatment_key in treatment_groups:
                treatment_groups[treatment_key]["quantity"] += 1
                treatment_groups[treatment_key]["total_price"] += visit.prix
            else:
                treatment_groups[treatment_key] = {
                    "treatment": visit.acte,
                    "tooth": visit.dent or "N/A",
                    "quantity": 1,
                    "unit_price": visit.prix,
                    "total_price": visit.prix
                }
        
        # Convert to list and sort by treatment name
        treatments = list(treatment_groups.values())
        treatments.sort(key=lambda x: x["treatment"])
        
        return treatments
    
    def create_invoice_data(self, visit_ids: List[int]) -> Dict:
        """Create complete invoice data structure"""
        visits = self.get_visits_for_invoice(visit_ids)
        if not visits:
            raise ValueError("No visits found for the provided IDs")
        
        patient = self.get_patient_by_visit_ids(visit_ids)
        if not patient:
            raise ValueError("Patient not found for the provided visits")
        
        treatments = self.aggregate_treatments(visits)
        totals = self.calculate_invoice_totals(visits)
        
        invoice_data = {
            "invoice_number": self.generate_invoice_number(),
            "invoice_date": datetime.now().strftime("%d/%m/%Y"),
            "patient": {
                "name": patient.full_name,
                "phone": patient.telephone or "",
                "national_id": patient.numero_carte_national or "",
                "assurance": patient.assurance or ""
            },
            "doctor": self.doctor_info,
            "treatments": treatments,
            "totals": totals,
            "visits": [
                {
                    "date": visit.date.strftime("%d/%m/%Y") if visit.date else "",
                    "treatment": visit.acte or "",
                    "tooth": visit.dent or "",
                    "price": visit.prix or 0.0
                }
                for visit in visits
            ]
        }
        
        return invoice_data
    
    def generate_invoice_number(self) -> str:
        """Generate a unique invoice number"""
        timestamp = datetime.now().strftime("%Y%m%d%H%M%S")
        return f"INV-{timestamp}"
    
    def create_word_invoice(self, invoice_data: Dict, output_path: str = None) -> str:
        """Create Word invoice document from template"""
        if not DOCX_AVAILABLE:
            raise ImportError("python-docx library is required for invoice generation")
        
        # Create new document
        doc = Document()
        
        # Set page margins
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(0.5)
            section.bottom_margin = Inches(0.5)
            section.left_margin = Inches(0.5)
            section.right_margin = Inches(0.5)
        
        # Add doctor header
        self.add_doctor_header(doc, invoice_data["doctor"])
        
        # Add invoice title
        self.add_invoice_title(doc)
        
        # Add patient information
        self.add_patient_info(doc, invoice_data["patient"])
        
        # Add totals section (simple format like the template)
        self.add_totals_section(doc, invoice_data["totals"])
        
        # Add footer
        self.add_footer(doc, invoice_data["doctor"])
        
        # Save document
        if not output_path:
            output_path = os.path.join(
                os.path.expanduser("~/.dentistedb/invoices"),
                f"invoice_{invoice_data['invoice_number']}.docx"
            )
        
        os.makedirs(os.path.dirname(output_path), exist_ok=True)
        doc.save(output_path)
        
        return output_path
    
    def add_doctor_header(self, doc: 'Document', doctor_info: Dict):
        """Add doctor information header"""
        # Doctor name and title
        name_para = doc.add_paragraph()
        name_run = name_para.add_run(doctor_info["name"])
        name_run.bold = True
        name_run.font.size = Inches(0.18)
        
        title_para = doc.add_paragraph()
        title_run = title_para.add_run(doctor_info["title"])
        title_run.font.size = Inches(0.14)
        
        # Address and contact
        address_para = doc.add_paragraph()
        address_para.add_run(doctor_info["address"])
        
        phone_para = doc.add_paragraph()
        phone_para.add_run(f"Tél : {doctor_info['phone']}")
        
        inpe_para = doc.add_paragraph()
        inpe_para.add_run(f"INPE {doctor_info['inpe']}")
        
        # Add spacing
        doc.add_paragraph()
    
    def add_invoice_title(self, doc: 'Document'):
        """Add invoice title"""
        title_para = doc.add_paragraph()
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title_para.add_run("Note d'honoraire")
        title_run.bold = True
        title_run.font.size = Inches(0.2)
        
        doc.add_paragraph()
    
    def add_patient_info(self, doc: 'Document', patient_info: Dict):
        """Add patient information section"""
        # Date
        date_para = doc.add_paragraph()
        date_para.add_run(f"Rabat le: {datetime.now().strftime('%d/%m/%Y')}")
        
        # Patient info
        patient_para = doc.add_paragraph()
        patient_para.add_run("Soins dentaires effectués pour :")
        
        name_para = doc.add_paragraph()
        name_run = name_para.add_run(patient_info["name"])
        name_run.bold = True
        
        if patient_info["phone"]:
            phone_para = doc.add_paragraph()
            phone_para.add_run(f"Téléphone: {patient_info['phone']}")
        
        if patient_info["national_id"]:
            id_para = doc.add_paragraph()
            id_para.add_run(f"Carte Nationale: {patient_info['national_id']}")
        
        if patient_info["assurance"]:
            assurance_para = doc.add_paragraph()
            assurance_para.add_run(f"Assurance: {patient_info['assurance']}")
        
        doc.add_paragraph()
    

    
    def add_totals_section(self, doc: 'Document', totals: Dict):
        """Add totals section"""
        # Simple total amount (like the template)
        total_para = doc.add_paragraph()
        total_para.add_run(f"Montant total: {totals['total']:.2f} DH")
        
        # Final statement
        final_para = doc.add_paragraph()
        final_para.add_run("Arrêtée la présente facture à la somme de")
        final_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        amount_para = doc.add_paragraph()
        amount_run = amount_para.add_run(f"{totals['total']:.2f} DH")
        amount_run.bold = True
        amount_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        payment_para = doc.add_paragraph()
        payment_para.add_run("Valeur en votre aimable règlement")
        payment_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph()
    
    def add_footer(self, doc: 'Document', doctor_info: Dict):
        """Add footer with doctor signature and details"""
        # Signature line
        signature_para = doc.add_paragraph()
        signature_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        signature_para.add_run(doctor_info["name"])
        
        # Tax information
        tax_para = doc.add_paragraph()
        tax_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        tax_text = (f"IF {doctor_info['if']} / Taxe professionnelle {doctor_info['taxe_professionnelle']} / "
                   f"Affiliation CNSS {doctor_info['cnss']} / ICE {doctor_info['ice']}")
        tax_para.add_run(tax_text)
    
    def convert_to_pdf(self, word_path: str, pdf_path: str = None) -> str:
        """Convert Word document to PDF (placeholder for future implementation)"""
        # This would require additional libraries like python-docx2pdf or similar
        # For now, return the Word document path
        if not pdf_path:
            pdf_path = word_path.replace('.docx', '.pdf')
        
        # Placeholder - in a real implementation, you would convert the Word doc to PDF
        print(f"PDF conversion not implemented. Word document saved at: {word_path}")
        return word_path 